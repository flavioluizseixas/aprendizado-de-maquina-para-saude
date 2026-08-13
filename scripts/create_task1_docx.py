"""Gera o enunciado em Word da Tarefa 1 da disciplina."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "tarefas" / "Tarefa_1_Estatistica_Descritiva.docx"

NAVY = "153A5B"
TEAL = "087E8B"
LIGHT_TEAL = "E8F4F5"
PALE_BLUE = "EDF3F8"
ORANGE = "E68A3F"
LIGHT_ORANGE = "FFF3E8"
GRAY = "5D6975"
LIGHT_GRAY = "F3F5F7"
WHITE = "FFFFFF"
BLACK = "1E2933"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D8E0E7", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])


def add_hyperlink(paragraph, text: str, url: str, color: str = TEAL):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    new_run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    run_properties.append(run_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)
    new_run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_text(paragraph, text: str, *, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def add_bullet(document, text: str, level: int = 0, bold_prefix: str | None = None):
    paragraph = document.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    if bold_prefix and text.startswith(bold_prefix):
        add_text(paragraph, bold_prefix, bold=True)
        add_text(paragraph, text[len(bold_prefix):])
    else:
        add_text(paragraph, text)
    return paragraph


def add_numbered_item(document, number: str, title: str, points: str, intro: str):
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(1.15)
    table.columns[1].width = Cm(15.8)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left, right = table.rows[0].cells
    left.width = Cm(1.15)
    right.width = Cm(15.8)
    for cell in (left, right):
        set_cell_border(cell, color="C9D8E3")
        set_cell_margins(cell, top=90, start=130, bottom=90, end=130)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(left, TEAL)
    set_cell_shading(right, PALE_BLUE)
    p_left = left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p_left, number, bold=True, color=WHITE, size=14)
    p_right = right.paragraphs[0]
    p_right.paragraph_format.space_after = Pt(2)
    add_text(p_right, title, bold=True, color=NAVY, size=11.5)
    add_text(p_right, f"  ·  {points}", bold=True, color=ORANGE, size=9)
    p_intro = document.add_paragraph(intro)
    p_intro.paragraph_format.space_after = Pt(4)
    return table


def add_deliverable(document, text: str):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_ORANGE)
    set_cell_border(cell, color="F2BE91")
    set_cell_margins(cell, top=90, start=140, bottom=90, end=140)
    p = cell.paragraphs[0]
    add_text(p, "ENTREGUE  ", bold=True, color=ORANGE, size=8.5)
    add_text(p, text, color=BLACK, size=9)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_callout(document, title: str, body: str, fill=LIGHT_TEAL, border="A7D6DA"):
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=border, size="8")
    set_cell_margins(cell, top=140, start=170, bottom=140, end=170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    add_text(p, title.upper(), bold=True, color=TEAL, size=8.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_text(p2, body, color=BLACK, size=9.5)
    return table


def style_document(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        styles[style_name].font.name = "Aptos Display"

    styles["Title"].font.size = Pt(30)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor.from_string(NAVY)
    styles["Title"].paragraph_format.space_after = Pt(8)

    styles["Subtitle"].font.size = Pt(14)
    styles["Subtitle"].font.color.rgb = RGBColor.from_string(TEAL)
    styles["Subtitle"].paragraph_format.space_after = Pt(18)

    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor.from_string(NAVY)
    styles["Heading 1"].paragraph_format.space_before = Pt(14)
    styles["Heading 1"].paragraph_format.space_after = Pt(6)
    styles["Heading 1"].paragraph_format.keep_with_next = True

    styles["Heading 2"].font.size = Pt(12.5)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor.from_string(TEAL)
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(4)
    styles["Heading 2"].paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Bullet 2", "List Number"):
        styles[list_style].font.name = "Aptos"
        styles[list_style].font.size = Pt(9.7)
        styles[list_style].paragraph_format.space_after = Pt(3)


def configure_section(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.65)


def add_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(p, "APRENDIZADO DE MÁQUINA PARA SAÚDE", bold=True, color=TEAL, size=8)
    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.8))
    table.columns[0].width = Cm(13.4)
    table.columns[1].width = Cm(3.4)
    p_left = table.cell(0, 0).paragraphs[0]
    add_text(p_left, "Tarefa 1  •  Estatística descritiva", color=GRAY, size=8)
    add_page_number(table.cell(0, 1).paragraphs[0])


def build_document() -> Document:
    document = Document()
    style_document(document)
    section = document.sections[0]
    configure_section(section)
    add_header_footer(section)

    # Faixa superior da capa.
    band = document.add_table(rows=1, cols=1)
    cell = band.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_border(cell, color=NAVY)
    set_cell_margins(cell, top=85, start=170, bottom=85, end=170)
    p = cell.paragraphs[0]
    add_text(p, "TAREFA 1", bold=True, color=WHITE, size=11)
    add_text(p, "   |   ATIVIDADE INDIVIDUAL   •   10,0 PONTOS   •   60–75 MIN", color=WHITE, size=9)

    document.add_paragraph().paragraph_format.space_after = Pt(12)
    title = document.add_paragraph(style="Title")
    title.add_run("Conhecendo indicadores\nde saúde")
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run("Estatística descritiva, visualização e associação")

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    add_text(p, "Pergunta orientadora  ", bold=True, color=ORANGE, size=9)
    add_text(
        p,
        "Como se distribuem indicadores autorrelatados de saúde nesta amostra, e quais conclusões eles não sustentam?",
        italic=True,
        color=NAVY,
        size=12,
    )

    info = document.add_table(rows=2, cols=3)
    info.autofit = False
    labels = [
        ("ALUNO(A)", "Nome completo"),
        ("TURMA", "________________"),
        ("DATA", "____ / ____ / ______"),
        ("PROFESSOR(A)", "____________________________"),
        ("ARQUIVO", "nome_sobrenome_tarefa1.ipynb"),
        ("VERSÃO", "1.0 • agosto/2026"),
    ]
    for i, cell in enumerate([c for row in info.rows for c in row.cells]):
        set_cell_shading(cell, LIGHT_GRAY if i % 2 == 0 else WHITE)
        set_cell_border(cell)
        set_cell_margins(cell, top=120, start=130, bottom=120, end=130)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_cell = cell.paragraphs[0]
        p_cell.paragraph_format.space_after = Pt(2)
        add_text(p_cell, labels[i][0], bold=True, color=TEAL, size=7.5)
        p_value = cell.add_paragraph()
        p_value.paragraph_format.space_after = Pt(0)
        add_text(p_value, labels[i][1], color=BLACK, size=9)

    document.add_paragraph().paragraph_format.space_after = Pt(5)
    add_callout(
        document,
        "Contexto",
        "Você analisará o conjunto CDC Diabetes Health Indicators (UCI 891), derivado do BRFSS. "
        "A base reúne indicadores de saúde e comportamentos autorrelatados. Seu objetivo é descrever a "
        "amostra, escolher representações coerentes com cada tipo de variável e comunicar associações sem "
        "transformá-las em conclusões causais ou clínicas.",
    )

    document.add_heading("O que você deverá demonstrar", level=1)
    objectives = [
        "obter e inspecionar uma base pública de forma reproduzível;",
        "distinguir variáveis numéricas, binárias e ordinais;",
        "resumir distribuições com medidas adequadas, incluindo mediana e intervalo interquartil;",
        "selecionar gráficos coerentes com o tipo de variável;",
        "comparar variáveis usando percentuais condicionais, testes e tamanhos de efeito;",
        "diferenciar descrição, associação, causalidade e utilidade clínica.",
    ]
    for objective in objectives:
        add_bullet(document, objective)

    document.add_heading("Base, ambiente e regras de comparabilidade", level=1)
    p = document.add_paragraph()
    add_text(p, "Fonte oficial: ", bold=True)
    add_hyperlink(
        p,
        "CDC Diabetes Health Indicators — UCI Machine Learning Repository (conjunto 891)",
        "https://archive.ics.uci.edu/dataset/891/cdc+diabetes+health+indicators",
    )
    add_bullet(document, "Use Python em Jupyter ou Google Colab, com pandas, matplotlib, seaborn, scipy, scikit-learn e statsmodels.")
    add_bullet(document, "Una os atributos ao alvo Diabetes_binary e não faça imputação nem remoção de linhas antes da inspeção.")
    add_bullet(document, "Use uma amostra estratificada de 20.000 registros, random_state = 42. Essa configuração torna seus números comparáveis ao gabarito.")
    add_bullet(document, "Use os rótulos: Diabetes_binary: 0 = sem diabetes; 1 = pré-diabetes/diabetes. HighBP: 0 = sem pressão alta; 1 = com pressão alta.")
    add_bullet(document, "Para GenHlth, preserve a ordem: 1 = excelente; 2 = muito boa; 3 = boa; 4 = regular; 5 = ruim.")
    add_bullet(document, "Os códigos de variáveis categóricas são rótulos; não os interprete como medidas contínuas.")

    add_callout(
        document,
        "Importante",
        "A resposta de referência é o notebook 01_estatistica_descritiva.ipynb e deve ser disponibilizada "
        "pelo docente somente após a realização da tarefa. O enunciado não apresenta resultados numéricos.",
        fill=LIGHT_ORANGE,
        border="F2BE91",
    )

    document.add_page_break()
    document.add_heading("Enunciado", level=1)
    p = document.add_paragraph(
        "Construa um notebook executável do início ao fim. Organize cada item com uma célula Markdown que "
        "explique a pergunta, seguida do código, da saída e de uma interpretação breve em linguagem própria."
    )
    p.paragraph_format.space_after = Pt(10)

    add_numbered_item(
        document,
        "1",
        "Obtenção dos dados e reprodutibilidade",
        "0,5 ponto",
        "Carregue o conjunto UCI 891, una atributos e alvo e produza a amostra definida nas regras de comparabilidade.",
    )
    add_bullet(document, "Registre o nome da base, o identificador UCI, a variável-alvo, o tamanho efetivo da amostra e a semente.")
    add_bullet(document, "Classifique as colunas em três grupos: binárias, ordinais e numéricas.")
    add_deliverable(document, "Código reproduzível + metadados da base + lista de variáveis por tipo.")

    add_numbered_item(
        document,
        "2",
        "Inspeção e qualidade aparente",
        "1,0 ponto",
        "Faça uma leitura inicial da tabela antes de calcular qualquer associação.",
    )
    add_bullet(document, "Mostre as cinco primeiras linhas, a dimensão da tabela e os tipos das colunas.")
    add_bullet(document, "Conte valores ausentes por coluna e linhas duplicadas.")
    add_bullet(document, "Calcule a distribuição percentual de Diabetes_binary e represente-a em gráfico de barras.")
    add_bullet(document, "Em duas ou três frases, diga o que essa proporção descreve e o que ela não mede.")
    add_deliverable(document, "Saídas da inspeção + tabela percentual + gráfico do alvo + interpretação.")

    add_numbered_item(
        document,
        "3",
        "Tabela de estatística descritiva",
        "1,0 ponto",
        "Resuma BMI, Age, GenHlth, PhysHlth, MentHlth, HighBP e Diabetes_binary.",
    )
    add_bullet(document, "Apresente n, média, desvio-padrão, Q1, mediana (Q2), Q3 e IQR = Q3 − Q1.")
    add_bullet(document, "Explique por que mediana e IQR merecem atenção em distribuições assimétricas ou com valores extremos.")
    add_bullet(document, "Não atribua significado quantitativo indevido aos códigos ordinais ou binários.")
    add_deliverable(document, "Uma tabela arredondada a duas casas decimais + interpretação de até cinco linhas.")

    add_numbered_item(
        document,
        "4",
        "Visualizações univariadas",
        "1,5 ponto",
        "Escolha a representação apropriada para uma variável numérica, uma binária e uma ordinal.",
    )
    add_bullet(document, "BMI: apresente a distribuição e uma visualização que evidencie quartis e valores extremos.")
    add_bullet(document, "HighBP: apresente categorias legíveis e percentuais.")
    add_bullet(document, "GenHlth: preserve a ordem natural das cinco categorias e use os rótulos em português.")
    add_bullet(document, "Inclua título, nomes dos eixos, unidade quando aplicável e tamanho da amostra.")
    add_deliverable(document, "Três análises visuais — cada uma acompanhada de uma conclusão descritiva curta.")

    add_numbered_item(
        document,
        "5",
        "Avaliação visual conjunta",
        "1,0 ponto",
        "Investigue conjuntamente BMI, PhysHlth e MentHlth, diferenciando os registros pela categoria de Diabetes_binary.",
    )
    add_bullet(document, "Construa um pairplot apenas com uma amostra estratificada de até 1.000 registros, usando random_state = 42.")
    add_bullet(document, "Explique por que a base completa não é usada nesse gráfico e como a estratificação ajuda na comparação.")
    add_bullet(document, "Comente a sobreposição entre os grupos sem afirmar capacidade diagnóstica ou causalidade.")
    add_deliverable(document, "Pairplot + justificativa da amostragem + leitura visual em até cinco linhas.")

    document.add_page_break()
    add_numbered_item(
        document,
        "6",
        "Duas variáveis categóricas: GenHlth × Diabetes_binary",
        "2,0 pontos",
        "Avalie a associação entre saúde geral autorrelatada e o indicador binário de diabetes.",
    )
    add_bullet(document, "Monte a tabela de contingência com contagens e a tabela de percentuais dentro de cada categoria de GenHlth.")
    add_bullet(document, "Visualize os percentuais com: (a) heatmap anotado; (b) barras 100% empilhadas; (c) gráfico de mosaico.")
    add_bullet(document, "Mostre também a proporção de pré-diabetes/diabetes em cada categoria de saúde geral.")
    add_bullet(document, "Calcule qui-quadrado, graus de liberdade, p-valor, menor frequência esperada e V de Cramér.")
    add_bullet(document, "Interprete conjuntamente direção/padrão, significância estatística e magnitude da associação. Explique por que um p-valor pequeno, em uma amostra grande, não basta.")
    add_deliverable(document, "Duas tabelas + quatro gráficos + medidas do teste + interpretação fundamentada.")

    add_numbered_item(
        document,
        "7",
        "Uma variável categórica e uma numérica: Diabetes_binary × BMI",
        "2,0 pontos",
        "Compare a distribuição de BMI entre os dois grupos do desfecho.",
    )
    add_bullet(document, "Por grupo, calcule n, média, mediana, desvio-padrão, Q1, Q3 e IQR.")
    add_bullet(document, "Produza um boxplot e um violin plot com quartis internos.")
    add_bullet(document, "Calcule a diferença das médias, a correlação ponto-bisserial e seu p-valor, η² = r² e a diferença padronizada de médias usando o desvio-padrão combinado.")
    add_bullet(document, "Interprete sinal e magnitude. Comente a sobreposição das distribuições e lembre que nenhuma medida controla confundimento.")
    add_deliverable(document, "Tabela por grupo + dois gráficos + cinco medidas de associação/efeito + interpretação.")

    add_numbered_item(
        document,
        "8",
        "Síntese, limitações e responsabilidade",
        "1,0 ponto",
        "Feche a análise com uma síntese clara, separando evidência observada de inferências não sustentadas.",
    )
    add_bullet(document, "Escreva três aprendizados principais obtidos nos itens anteriores.")
    add_bullet(document, "Discuta ao menos: autorrelato e possível erro de memória/classificação; amostragem e representatividade; desbalanceamento; associação versus causalidade.")
    add_bullet(document, "Declare por que os achados não devem ser usados isoladamente para diagnóstico, prognóstico, tratamento ou decisão de saúde pública.")
    add_bullet(document, "Registre as versões das bibliotecas utilizadas.")
    add_deliverable(document, "Síntese de 150–250 palavras + versões do ambiente.")

    document.add_heading("Formato de entrega", level=1)
    delivery = [
        "Entregue um arquivo .ipynb executado do início ao fim, salvo orientação docente em contrário.",
        "Use o padrão nome_sobrenome_tarefa1.ipynb.",
        "Todas as tabelas e figuras devem aparecer na saída; não entregue apenas código.",
        "Textos interpretativos devem ser autorais, objetivos e coerentes com os resultados exibidos.",
        "Informe fonte dos dados, parâmetros de amostragem e versões do ambiente.",
    ]
    for item in delivery:
        add_bullet(document, item)

    document.add_heading("Critérios de avaliação", level=1)
    rows = [
        ("1", "Dados e reprodutibilidade", "0,5"),
        ("2", "Inspeção e qualidade aparente", "1,0"),
        ("3", "Estatística descritiva", "1,0"),
        ("4", "Visualizações univariadas", "1,5"),
        ("5", "Avaliação visual conjunta", "1,0"),
        ("6", "Comparação categórica–categórica", "2,0"),
        ("7", "Comparação categórica–numérica", "2,0"),
        ("8", "Síntese, limitações e ambiente", "1,0"),
        ("", "TOTAL", "10,0"),
    ]
    rubric = document.add_table(rows=1, cols=3)
    rubric.autofit = False
    widths = [Cm(1.5), Cm(12.7), Cm(2.4)]
    for i, width in enumerate(widths):
        rubric.columns[i].width = width
    headers = ["ITEM", "CRITÉRIO", "PONTOS"]
    for i, cell in enumerate(rubric.rows[0].cells):
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, color=NAVY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
        add_text(p, headers[i], bold=True, color=WHITE, size=8)
    set_repeat_table_header(rubric.rows[0])
    for row_index, row_data in enumerate(rows):
        cells = rubric.add_row().cells
        for col_index, value in enumerate(row_data):
            set_cell_shading(cells[col_index], LIGHT_GRAY if row_index % 2 == 0 else WHITE)
            set_cell_border(cells[col_index])
            set_cell_margins(cells[col_index], top=75, start=120, bottom=75, end=120)
            p = cells[col_index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_index != 1 else WD_ALIGN_PARAGRAPH.LEFT
            add_text(p, value, bold=row_index == len(rows) - 1, size=8.7)

    document.add_page_break()
    document.add_heading("Roteiro de autocorreção", level=1)
    p = document.add_paragraph(
        "Use esta seção somente depois que o docente liberar a resposta de referência. Compare raciocínio, "
        "parâmetros e resultados — não apenas a aparência dos gráficos. Pequenas diferenças de versão podem "
        "alterar formatação, mas não devem mudar as conclusões centrais."
    )
    p.paragraph_format.space_after = Pt(10)

    checks = [
        ("Reprodutibilidade", "A amostra, a semente e a estratificação são idênticas às da referência?"),
        ("Inspeção", "Dimensão, tipos, ausências, duplicatas e distribuição do alvo coincidem?"),
        ("Escolha gráfica", "Cada gráfico respeita o caráter numérico, binário ou ordinal da variável?"),
        ("Percentuais", "Na comparação categórica, os percentuais foram calculados dentro de cada categoria de GenHlth?"),
        ("Tamanhos de efeito", "V de Cramér, r ponto-bisserial, η² e diferença padronizada têm sinal/magnitude compatíveis?"),
        ("Interpretação", "A conclusão considera magnitude e sobreposição, em vez de se apoiar apenas no p-valor?"),
        ("Limites", "O texto evita causalidade, diagnóstico e generalização indevida?"),
        ("Execução", "O notebook roda de cima para baixo sem depender de estado oculto?"),
    ]
    checklist = document.add_table(rows=1, cols=3)
    checklist.autofit = False
    checklist.columns[0].width = Cm(1.2)
    checklist.columns[1].width = Cm(4.2)
    checklist.columns[2].width = Cm(11.1)
    for i, value in enumerate(("OK", "ASPECTO", "PERGUNTA PARA CONFERÊNCIA")):
        cell = checklist.rows[0].cells[i]
        set_cell_shading(cell, TEAL)
        set_cell_border(cell, color=TEAL)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        add_text(p, value, bold=True, color=WHITE, size=8)
    set_repeat_table_header(checklist.rows[0])
    for row_index, (aspect, question) in enumerate(checks):
        cells = checklist.add_row().cells
        values = ("☐", aspect, question)
        for col_index, value in enumerate(values):
            set_cell_shading(cells[col_index], PALE_BLUE if row_index % 2 == 0 else WHITE)
            set_cell_border(cells[col_index])
            set_cell_margins(cells[col_index], top=105, start=120, bottom=105, end=120)
            p = cells[col_index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_text(p, value, bold=col_index == 1, size=8.8)

    document.add_heading("Registro da comparação", level=2)
    for prompt in (
        "Uma diferença que encontrei entre minha solução e a referência:",
        "O que corrigi ou compreendi após a comparação:",
        "Uma dúvida que permaneceu:",
    ):
        p = document.add_paragraph()
        add_text(p, prompt, bold=True, color=NAVY, size=9.5)
        for _ in range(2):
            line = document.add_paragraph("________________________________________________________________________________")
            line.paragraph_format.space_after = Pt(2)
            for run in line.runs:
                run.font.color.rgb = RGBColor.from_string("B7C2CC")

    add_callout(
        document,
        "Uso responsável",
        "Este material tem finalidade exclusivamente educacional. Resultados descritivos e associações em "
        "dados observacionais e autorrelatados não substituem avaliação clínica nem validação apropriada.",
    )

    return document


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.core_properties.title = "Tarefa 1 — Estatística descritiva e visualização"
    document.core_properties.subject = "Aprendizado de Máquina para Saúde"
    document.core_properties.author = "Flavio Luiz Seixas"
    document.core_properties.keywords = "saúde, estatística descritiva, visualização, CDC Diabetes, tarefa"
    document.core_properties.comments = (
        "Enunciado alinhado ao notebook de referência 01_estatistica_descritiva.ipynb."
    )
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
